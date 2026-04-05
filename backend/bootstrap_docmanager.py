"""
RAGraph — Document Management System
Implements all 10 features in one bootstrap.
Run: python bootstrap_docmanager.py
from A:\Projects\RAGraph\backend\
"""
import os

BASE = os.path.dirname(os.path.abspath(__file__))
FRONTEND = os.path.join(os.path.dirname(BASE), "frontend", "src")

def w(path, content, base=None):
    root = base or BASE
    full = os.path.join(root, path.replace("/", os.sep))
    os.makedirs(os.path.dirname(full), exist_ok=True)
    with open(full, "w", encoding="utf-8") as f:
        f.write(content.lstrip("\n"))
    print(f"  wrote: {path}")

print("\n" + "="*55)
print("  RAGraph — Document Management System")
print("="*55)

# ═══════════════════════════════════════════════════════════
# BACKEND: app/api/documents.py — full document management
# ═══════════════════════════════════════════════════════════
w("app/api/documents.py", '''
from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import PlainTextResponse
from loguru import logger
import uuid, os, time
from app.config import settings
from app.models.search import UploadResponse
from app.models.documents import DocumentMetadata, DocumentStatus

router = APIRouter(prefix="/api/documents", tags=["documents"])

ALLOWED_EXT = {"pdf", "docx", "pptx", "txt", "md"}
MAX_FILE_SIZE = 50 * 1024 * 1024

# In-memory store (guest session — resets on server restart)
# Structure: { doc_id: { meta, number } }
_docs: dict[str, dict] = {}
_counter = 0  # sequential numbering


def _next_number() -> int:
    """Assign next available sequential number."""
    global _counter
    _counter += 1
    return _counter


def _renumber():
    """Renumber all docs sequentially after deletion."""
    sorted_docs = sorted(_docs.values(), key=lambda d: d["number"])
    for i, doc in enumerate(sorted_docs, 1):
        doc["number"] = i


@router.post("/upload", response_model=UploadResponse)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
):
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(status_code=415,
            detail=f"Unsupported .{ext}. Allowed: {ALLOWED_EXT}")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413,
            detail=f"File too large ({len(content)/1024/1024:.1f}MB). Max 50MB.")

    doc_id = str(uuid.uuid4())
    safe_filename = f"{doc_id}.{ext}"
    storage_path = os.path.join(settings.local_storage_path, safe_filename)
    os.makedirs(settings.local_storage_path, exist_ok=True)

    with open(storage_path, "wb") as f:
        f.write(content)

    original_name = file.filename or safe_filename
    number = _next_number()

    metadata = DocumentMetadata(
        id=doc_id,
        filename=safe_filename,
        original_filename=original_name,
        mime_type=file.content_type or f"application/{ext}",
        storage_path=storage_path,
        status=DocumentStatus.QUEUED,
    )
    _docs[doc_id] = {"meta": metadata, "number": number}

    logger.info(f"[Doc {number}] Uploaded: {original_name} ({len(content):,} bytes)")
    background_tasks.add_task(_run_ingestion, metadata)

    return UploadResponse(
        doc_id=doc_id,
        filename=original_name,
        status="queued",
        message=f"[Doc {number}] {original_name} queued for processing.",
    )


@router.get("/{doc_id}/status")
async def get_status(doc_id: str):
    entry = _docs.get(doc_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Document not found")
    meta = entry["meta"]
    return {
        "doc_id": doc_id,
        "number": entry["number"],
        "filename": meta.original_filename,
        "status": meta.status,
        "page_count": meta.page_count,
        "node_count": meta.node_count,
        "image_count": meta.image_count,
        "h1_count": meta.h1_count,
        "paragraph_count": meta.paragraph_count,
        "error": meta.error_message,
        "ingested_at": meta.ingested_at.isoformat() if meta.ingested_at else None,
    }


@router.get("/")
async def list_documents():
    """List all documents with sequential numbers."""
    docs = []
    for doc_id, entry in _docs.items():
        meta = entry["meta"]
        docs.append({
            "doc_id": doc_id,
            "number": entry["number"],
            "filename": meta.original_filename,
            "status": meta.status,
            "node_count": meta.node_count,
            "image_count": meta.image_count,
            "page_count": meta.page_count,
            "size_bytes": os.path.getsize(meta.storage_path)
              if os.path.exists(meta.storage_path) else 0,
        })
    # Sort by number
    docs.sort(key=lambda d: d["number"])
    return {"documents": docs, "total": len(docs)}


@router.delete("/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document, remove from Qdrant, renumber remaining."""
    entry = _docs.get(doc_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Document not found")

    meta = entry["meta"]
    number = entry["number"]

    # Remove from Qdrant
    try:
        from app.services.qdrant_service import qdrant_service
        await qdrant_service.delete_by_doc(doc_id)
        logger.info(f"[Doc {number}] Removed from Qdrant: {meta.original_filename}")
    except Exception as e:
        logger.warning(f"Qdrant delete failed: {e}")

    # Remove file from disk
    try:
        if os.path.exists(meta.storage_path):
            os.remove(meta.storage_path)
    except Exception as e:
        logger.warning(f"File delete failed: {e}")

    # Remove from in-memory store
    del _docs[doc_id]

    # Renumber remaining docs
    _renumber()

    logger.info(f"[Doc {number}] Deleted: {meta.original_filename}. "
                f"Remaining: {len(_docs)}")

    return {
        "deleted": doc_id,
        "filename": meta.original_filename,
        "documents_remaining": len(_docs),
    }


@router.get("/{doc_id}/content")
async def get_document_content(doc_id: str):
    """Return document text content for the viewer."""
    entry = _docs.get(doc_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Document not found")

    meta = entry["meta"]
    if not os.path.exists(meta.storage_path):
        raise HTTPException(status_code=404, detail="File not found on disk")

    ext = meta.filename.rsplit(".", 1)[-1].lower()

    try:
        if ext == "txt" or ext == "md":
            with open(meta.storage_path, "r", encoding="utf-8", errors="replace") as f:
                return PlainTextResponse(f.read())

        elif ext == "pdf":
            import fitz
            doc = fitz.open(meta.storage_path)
            pages = []
            for i, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    pages.append(f"--- Page {i+1} ---\\n{text}")
            doc.close()
            return PlainTextResponse("\\n\\n".join(pages[:20]))  # first 20 pages

        elif ext == "docx":
            from docx import Document
            doc = Document(meta.storage_path)
            text = "\\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return PlainTextResponse(text)

        else:
            return PlainTextResponse(f"Preview not available for .{ext} files")

    except Exception as e:
        return PlainTextResponse(f"Could not read file: {e}")


def get_doc_number_map() -> dict[str, int]:
    """Returns {doc_id: number} for citation mapping."""
    return {doc_id: entry["number"] for doc_id, entry in _docs.items()}


def get_doc_name_map() -> dict[str, str]:
    """Returns {doc_id: original_filename} for citation display."""
    return {doc_id: entry["meta"].original_filename for doc_id, entry in _docs.items()}


async def _run_ingestion(metadata: DocumentMetadata) -> None:
    from app.core.ingestion.pipeline import run_ingestion
    updated = await run_ingestion(metadata)
    if metadata.id in _docs:
        _docs[metadata.id]["meta"] = updated
''')

# ═══════════════════════════════════════════════════════════
# BACKEND: app/api/search.py — citation mapping
# ═══════════════════════════════════════════════════════════
w("app/api/search.py", '''
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from loguru import logger
import json, time, asyncio
from app.models.search import (
    SearchRequest, SearchResponse,
    StreamChunk, StreamChunkType,
    SourceItem, ImageItem,
)
from app.config import settings
from app.services.qdrant_service import qdrant_service
from app.services.redis_service import redis_service

router = APIRouter(prefix="/api/search", tags=["search"])


@router.post("/stream")
async def stream_search(request: SearchRequest):
    logger.info(f"Search: \'{request.query[:60]}\' model={request.model} "
                f"hyde={request.use_hyde} has_image={bool(request.image)}")
    return StreamingResponse(
        _pipeline(request),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache",
                 "X-Accel-Buffering": "no",
                 "Connection": "keep-alive"},
    )


async def _pipeline(request: SearchRequest):
    start = time.time()
    answer_buffer = ""

    try:
        # Cache check (skip if empty chunks stored)
        if not request.image:
            cached = await redis_service.get_cached_result(
                request.query, request.model, request.focus
            )
            if cached and cached.get("chunk_count", 0) > 0:
                logger.info("Cache hit")
                async for chunk in _replay_cached(cached):
                    yield chunk
                return

        # Retrieve
        from app.core.retrieval.orchestrator import retrieve
        retrieval = await retrieve(
            query=request.query,
            use_hyde=request.use_hyde,
            use_dual_path=request.use_dual_path,
            image_base64=request.image,
            top_k=settings.top_k_final,
        )

        # Get doc number + name maps for citations
        from app.api.documents import get_doc_number_map, get_doc_name_map
        doc_numbers = get_doc_number_map()
        doc_names   = get_doc_name_map()

        # Build source items with doc number
        sources = _chunks_to_sources(retrieval.chunks, doc_numbers, doc_names)
        yield _sse(StreamChunkType.SOURCES,
                   json.dumps([s.model_dump() for s in sources]))

        images = _to_image_items(retrieval.images)
        yield _sse(StreamChunkType.IMAGES,
                   json.dumps([i.model_dump() for i in images]))

        # Build context with doc-number citations
        from app.core.generation.context_builder import build_context
        context = build_context(
            query=request.query,
            chunks=retrieval.chunks,
            images=retrieval.images,
            conversation_history=request.conversation_history,
            image_base64=request.image,
            doc_numbers=doc_numbers,
            doc_names=doc_names,
        )

        # Stream LLM answer
        from app.core.generation.llm_client import (
            stream_answer, generate_related_questions
        )
        async for delta in stream_answer(
            context=context,
            model=request.model,
            image_base64=request.image,
        ):
            answer_buffer += delta
            yield _sse(StreamChunkType.TEXT, delta)

        # Citation map for display
        citation_map = _build_citation_map(retrieval.chunks, doc_numbers, doc_names)
        yield _sse(StreamChunkType.CITATIONS, json.dumps(citation_map))

        # Related questions
        headings = list({
            c.heading_path[-1] for c in retrieval.chunks if c.heading_path
        })
        related = await generate_related_questions(
            query=request.query,
            answer=answer_buffer,
            headings=headings,
        )
        yield _sse(StreamChunkType.RELATED, json.dumps(related))

        elapsed_ms = int((time.time() - start) * 1000)
        done_meta = {
            "time_ms": elapsed_ms,
            "tokens": context.total_tokens + len(answer_buffer) // 4,
            "model": request.model,
            "hyde_used": retrieval.hyde_used,
            "dual_path_fallback_used": retrieval.dual_path_fallback_used,
            "chunks_used": len(retrieval.chunks),
            "strategies_used": getattr(retrieval, "strategies_used", []),
        }
        yield _sse(StreamChunkType.DONE, json.dumps(done_meta))

        # Cache only if we got real chunks
        if not request.image and answer_buffer and len(retrieval.chunks) > 0:
            await redis_service.cache_result(
                request.query, request.model, request.focus,
                {
                    "sources":     [s.model_dump() for s in sources],
                    "images":      [i.model_dump() for i in images],
                    "answer":      answer_buffer,
                    "related":     related,
                    "meta":        done_meta,
                    "citation_map": citation_map,
                    "chunk_count": len(retrieval.chunks),
                },
            )

    except asyncio.CancelledError:
        logger.info("Client disconnected")
    except Exception as e:
        logger.error("Pipeline error: {}", repr(e))
        yield _sse(StreamChunkType.ERROR, str(e))


def _sse(chunk_type, content: str) -> str:
    return StreamChunk(type=chunk_type, content=content).to_sse()


async def _replay_cached(cached: dict):
    yield _sse(StreamChunkType.SOURCES, json.dumps(cached.get("sources", [])))
    yield _sse(StreamChunkType.IMAGES,  json.dumps(cached.get("images", [])))
    answer = cached.get("answer", "")
    for word in answer.split(" "):
        await asyncio.sleep(0.008)
        yield _sse(StreamChunkType.TEXT, word + " ")
    yield _sse(StreamChunkType.CITATIONS, json.dumps(cached.get("citation_map", {})))
    yield _sse(StreamChunkType.RELATED,  json.dumps(cached.get("related", [])))
    yield _sse(StreamChunkType.DONE,     json.dumps(cached.get("meta", {})))


def _chunks_to_sources(chunks, doc_numbers, doc_names) -> list[SourceItem]:
    sources = []
    for c in chunks:
        num = doc_numbers.get(c.doc_id, 0)
        name = doc_names.get(c.doc_id, "Unknown")
        label = f"[Doc {num}] {name}" if num else name
        sources.append(SourceItem(
            id=c.node_id,
            title=c.heading_path[-1] if c.heading_path else label,
            url=f"#chunk-{c.node_id[:8]}",
            domain=f"Doc {num}" if num else "ragraph",
            snippet=c.text[:200],
            relevance_score=round(c.score, 3),
            heading_path=c.heading_path,
            doc_id=c.doc_id,
        ))
    return sources


def _to_image_items(images) -> list[ImageItem]:
    return [
        ImageItem(
            id=img.image_id, url=img.storage_url,
            caption=img.caption,
            source_title=img.nearest_heading or "Document Image",
            source_url=img.storage_url,
            heading_path=img.heading_path,
            relevance_score=round(img.score, 3),
            width=img.width, height=img.height,
        )
        for img in images
    ]


def _build_citation_map(chunks, doc_numbers, doc_names) -> dict:
    """Build citation number -> doc info mapping for display."""
    seen = {}
    result = {}
    for i, chunk in enumerate(chunks, 1):
        doc_id = chunk.doc_id
        if doc_id not in seen:
            seen[doc_id] = True
            num = doc_numbers.get(doc_id, 0)
            name = doc_names.get(doc_id, "Unknown")
            result[str(i)] = {
                "doc_number": num,
                "filename": name,
                "label": f"Doc {num}" if num else "Unknown",
                "heading": chunk.heading_path[-1] if chunk.heading_path else "",
            }
    return result


@router.post("", response_model=SearchResponse)
async def search(request: SearchRequest):
    start = time.time()
    sources, images, answer, related = [], [], "", []
    meta: dict = {}
    async for raw in _pipeline(request):
        try:
            line = raw.strip()
            if not line.startswith("data: "):
                continue
            chunk = json.loads(line[6:])
            if chunk["type"] == "sources":
                sources = json.loads(chunk["content"])
            elif chunk["type"] == "images":
                images = json.loads(chunk["content"])
            elif chunk["type"] == "text":
                answer += chunk["content"]
            elif chunk["type"] == "related":
                related = json.loads(chunk["content"])
            elif chunk["type"] == "done":
                meta = json.loads(chunk["content"])
        except Exception:
            pass
    return SearchResponse(
        query=request.query, answer=answer,
        sources=[SourceItem(**s) for s in sources],
        images=[ImageItem(**i) for i in images],
        related_questions=related,
        model_used=meta.get("model", request.model),
        tokens_used=meta.get("tokens", 0),
        time_ms=int((time.time() - start) * 1000),
        hyde_used=meta.get("hyde_used", False),
        dual_path_fallback_used=meta.get("dual_path_fallback_used", False),
    )
''')

# ═══════════════════════════════════════════════════════════
# BACKEND: update StreamChunkType to include CITATIONS
# ═══════════════════════════════════════════════════════════
models_path = os.path.join(BASE, "app", "models", "search.py")
with open(models_path, "r") as f:
    models = f.read()
if "CITATIONS" not in models:
    models = models.replace(
        '    DONE    = "done"\n    ERROR   = "error"',
        '    DONE      = "done"\n    ERROR     = "error"\n    CITATIONS = "citations"'
    )
    with open(models_path, "w") as f:
        f.write(models)
    print("  app/models/search.py: CITATIONS chunk type added")

# ═══════════════════════════════════════════════════════════
# BACKEND: context_builder.py — doc-number citations
# ═══════════════════════════════════════════════════════════
w("app/core/generation/context_builder.py", '''
from __future__ import annotations
from dataclasses import dataclass, field
from typing import Optional
import tiktoken
from loguru import logger

from app.core.retrieval.tree_retriever import RetrievedChunk
from app.core.retrieval.image_retriever import RetrievedImage
from app.models.search import ConversationMessage

SYSTEM_PROMPT_TOKENS  = 400
CONTEXT_BUDGET_TOKENS = 8000
HISTORY_BUDGET_TOKENS = 1500
_enc = tiktoken.get_encoding("cl100k_base")

def _count(text: str) -> int:
    return len(_enc.encode(text))


@dataclass
class BuiltContext:
    system_prompt: str
    user_message: str
    citation_map: dict[int, RetrievedChunk]
    image_urls: list[str]
    total_tokens: int


SYSTEM_PROMPT = """You are RAGraph, an intelligent research assistant.
Answer questions based on the provided document context.

Rules:
1. Cite sources as [Doc N] where N is the document number shown in context.
   Example: "The perceptron is a linear classifier [Doc 1]."
2. If not found in context, say "I couldn\'t find this in the provided documents."
3. Be precise. Use markdown for clarity.
4. When referencing figures, mention them explicitly."""


def build_context(
    query: str,
    chunks: list[RetrievedChunk],
    images: list[RetrievedImage],
    conversation_history: list[ConversationMessage],
    image_base64: Optional[str] = None,
    doc_numbers: dict[str, int] = None,
    doc_names: dict[str, str] = None,
) -> BuiltContext:
    doc_numbers = doc_numbers or {}
    doc_names   = doc_names or {}
    citation_map: dict[int, RetrievedChunk] = {}
    context_parts: list[str] = []
    tokens_used = 0

    for i, chunk in enumerate(chunks, 1):
        doc_num  = doc_numbers.get(chunk.doc_id, 0)
        doc_name = doc_names.get(chunk.doc_id, "Unknown")
        heading  = " > ".join(chunk.heading_path) if chunk.heading_path else "Document"

        # Citation label uses Doc number
        label = f"[Doc {doc_num}]" if doc_num else f"[{doc_name}]"
        citation_text = (
            f"{label} **{heading}** (relevance: {chunk.score:.2f})\\n"
            f"{chunk.text}\\n"
        )
        chunk_tokens = _count(citation_text)
        if tokens_used + chunk_tokens > CONTEXT_BUDGET_TOKENS:
            break
        context_parts.append(citation_text)
        citation_map[i] = chunk
        tokens_used += chunk_tokens

    # Image context
    image_urls: list[str] = []
    img_parts: list[str] = []
    for img in images[:3]:
        label   = img.fig_label or "Image"
        heading = img.nearest_heading or " > ".join(img.heading_path)
        desc = f"[Image: {label}] Section: {heading}"
        if img.caption:
            desc += f" | Caption: {img.caption}"
        img_parts.append(desc)
        image_urls.append(img.storage_url)

    # Conversation history
    history_text = ""
    history_tokens = 0
    for msg in conversation_history[-6:]:
        line = f"{msg.role.upper()}: {msg.content}\\n"
        t = _count(line)
        if history_tokens + t > HISTORY_BUDGET_TOKENS:
            break
        history_text += line
        history_tokens += t

    # Assemble user message
    parts = []
    if history_text:
        parts.append(f"Previous conversation:\\n{history_text}")

    # Add doc number legend at top of context
    if doc_numbers:
        legend_lines = []
        seen = set()
        for chunk in chunks[:len(citation_map)]:
            if chunk.doc_id not in seen:
                seen.add(chunk.doc_id)
                num  = doc_numbers.get(chunk.doc_id, 0)
                name = doc_names.get(chunk.doc_id, "Unknown")
                if num:
                    legend_lines.append(f"[Doc {num}] = {name}")
        if legend_lines:
            parts.append("Document Legend:\\n" + "\\n".join(legend_lines))

    parts.append("Document context:\\n" + "\\n\\n".join(context_parts))
    if img_parts:
        parts.append("Available images:\\n" + "\\n".join(img_parts))
    parts.append(f"Question: {query}")

    user_message = "\\n\\n".join(parts)
    total_tokens = _count(SYSTEM_PROMPT) + _count(user_message)

    logger.debug(f"Context built: {len(citation_map)} citations, "
                 f"{len(image_urls)} images, {total_tokens} tokens")

    return BuiltContext(
        system_prompt=SYSTEM_PROMPT,
        user_message=user_message,
        citation_map=citation_map,
        image_urls=image_urls,
        total_tokens=total_tokens,
    )
''')

# ═══════════════════════════════════════════════════════════
# FRONTEND: src/lib/types.ts
# ═══════════════════════════════════════════════════════════
w("lib/types.ts", '''
export interface Source {
  id: string
  title: string
  url: string
  favicon?: string
  domain: string
  snippet: string
  relevance_score: number
  heading_path: string[]
  doc_id?: string
}

export interface RetrievedImage {
  id: string
  url: string
  caption?: string
  alt?: string
  source_title: string
  source_url: string
  heading_path: string[]
  relevance_score: number
  width?: number
  height?: number
}

export interface CitationEntry {
  doc_number: number
  filename: string
  label: string
  heading: string
}

export interface DocumentInfo {
  doc_id: string
  number: number
  filename: string
  status: \'queued\' | \'parsing\' | \'embedding\' | \'indexing\' | \'done\' | \'error\'
  node_count: number
  image_count: number
  page_count: number
  size_bytes: number
}

export interface Message {
  id: string
  role: \'user\' | \'assistant\'
  content: string
  sources?: Source[]
  images?: RetrievedImage[]
  related_questions?: string[]
  citation_map?: Record<string, CitationEntry>
  timestamp: Date
  isStreaming?: boolean
  meta?: {
    time_ms?: number
    tokens?: number
    hyde_used?: boolean
    dual_path_fallback_used?: boolean
    chunks_used?: number
    strategies_used?: string[]
  }
}

export interface Thread {
  id: string
  title: string
  messages: Message[]
  createdAt: Date
  updatedAt: Date
  model: string
  focus: FocusMode
}

export type FocusMode = \'all\' | \'academic\' | \'code\' | \'news\' | \'images\'

export const FOCUS_LABELS: Record<FocusMode, { label: string; icon: string }> = {
  all:      { label: \'All\',      icon: \'Globe\' },
  academic: { label: \'Academic\', icon: \'GraduationCap\' },
  code:     { label: \'Code\',     icon: \'Code2\' },
  news:     { label: \'News\',     icon: \'Newspaper\' },
  images:   { label: \'Images\',   icon: \'Image\' },
}

export type ModelOption = string

export const MODEL_LABELS: Record<string, string> = {
  \'gpt-4o\':                                        \'GPT-4o\',
  \'claude-3-5-sonnet\':                             \'Claude 3.5\',
  \'openrouter/free\':                               \'Auto Free\',
  \'meta-llama/llama-3.3-70b-instruct:free\':        \'Llama 3.3 (Free)\',
  \'mistralai/mistral-small-3.1-24b-instruct:free\': \'Mistral Small (Free)\',
  \'google/gemma-3-27b-it:free\':                    \'Gemma 3 (Free)\',
}

export interface SearchRequest {
  query: string
  model: string
  focus: FocusMode
  thread_id?: string
  conversation_history?: Array<{ role: \'user\' | \'assistant\'; content: string }>
  image?: string
  use_hyde?: boolean
  use_dual_path?: boolean
}

export interface UploadResponse {
  doc_id: string
  filename: string
  status: string
  message: string
}
''', base=FRONTEND)

# ═══════════════════════════════════════════════════════════
# FRONTEND: src/lib/api.ts
# ═══════════════════════════════════════════════════════════
w("lib/api.ts", '''
import { parseSSEStream } from \'./utils\'
import type { SearchRequest, Source, RetrievedImage, UploadResponse, DocumentInfo, CitationEntry } from \'./types\'

const BASE_URL = process.env.NEXT_PUBLIC_API_URL ?? \'http://localhost:8000\'

export async function healthCheck() {
  const res = await fetch(`${BASE_URL}/health`)
  if (!res.ok) throw new Error(\'Backend unreachable\')
  return res.json()
}

export interface StreamCallbacks {
  onText:      (delta: string) => void
  onSources:   (sources: Source[]) => void
  onImages:    (images: RetrievedImage[]) => void
  onCitations: (map: Record<string, CitationEntry>) => void
  onRelated:   (questions: string[]) => void
  onDone:      (meta: Record<string, unknown>) => void
  onError:     (error: string) => void
}

export async function streamSearch(
  request: SearchRequest,
  callbacks: StreamCallbacks,
  abortSignal?: AbortSignal,
): Promise<void> {
  const response = await fetch(`${BASE_URL}/api/search/stream`, {
    method: \'POST\',
    headers: { \'Content-Type\': \'application/json\' },
    body: JSON.stringify(request),
    signal: abortSignal,
  })
  if (!response.ok) {
    const err = await response.json().catch(() => ({ message: \'Request failed\' }))
    callbacks.onError(err.detail ?? err.message ?? \'Unknown error\')
    return
  }
  for await (const raw of parseSSEStream(response)) {
    try {
      const chunk = JSON.parse(raw) as { type: string; content: string }
      switch (chunk.type) {
        case \'text\':      callbacks.onText(chunk.content); break
        case \'sources\':   callbacks.onSources(JSON.parse(chunk.content)); break
        case \'images\':    callbacks.onImages(JSON.parse(chunk.content)); break
        case \'citations\': callbacks.onCitations(JSON.parse(chunk.content)); break
        case \'related\':   callbacks.onRelated(JSON.parse(chunk.content)); break
        case \'done\':      callbacks.onDone(JSON.parse(chunk.content)); break
        case \'error\':     callbacks.onError(chunk.content); break
      }
    } catch { /* skip malformed */ }
  }
}

export async function uploadDocument(
  file: File,
  onProgress?: (pct: number) => void,
): Promise<UploadResponse> {
  return new Promise((resolve, reject) => {
    const xhr = new XMLHttpRequest()
    xhr.open(\'POST\', `${BASE_URL}/api/documents/upload`)
    xhr.upload.onprogress = (e) => {
      if (onProgress && e.total) onProgress(Math.round((e.loaded / e.total) * 100))
    }
    xhr.onload = () => {
      if (xhr.status >= 200 && xhr.status < 300) resolve(JSON.parse(xhr.responseText))
      else reject(new Error(JSON.parse(xhr.responseText)?.detail ?? \'Upload failed\'))
    }
    xhr.onerror = () => reject(new Error(\'Network error\'))
    const form = new FormData()
    form.append(\'file\', file)
    xhr.send(form)
  })
}

export async function pollIngestionStatus(docId: string) {
  const res = await fetch(`${BASE_URL}/api/documents/${docId}/status`)
  if (!res.ok) throw new Error(\'Status check failed\')
  return res.json()
}

export async function listDocuments(): Promise<{ documents: DocumentInfo[]; total: number }> {
  const res = await fetch(`${BASE_URL}/api/documents/`)
  if (!res.ok) throw new Error(\'Failed to list documents\')
  return res.json()
}

export async function deleteDocument(docId: string) {
  const res = await fetch(`${BASE_URL}/api/documents/${docId}`, { method: \'DELETE\' })
  if (!res.ok) throw new Error(\'Delete failed\')
  return res.json()
}

export async function getDocumentContent(docId: string): Promise<string> {
  const res = await fetch(`${BASE_URL}/api/documents/${docId}/content`)
  if (!res.ok) throw new Error(\'Could not load document\')
  return res.text()
}

export async function clearCache() {
  await fetch(`${BASE_URL}/cache/clear`, { method: \'DELETE\' })
}

export function resolveImageUrl(storageUrl: string): string {
  if (!storageUrl) return \'\'
  if (storageUrl.startsWith(\'http\')) return storageUrl
  return storageUrl
}
''', base=FRONTEND)

# ═══════════════════════════════════════════════════════════
# FRONTEND: src/store/useSearchStore.ts
# ═══════════════════════════════════════════════════════════
w("store/useSearchStore.ts", '''
\'use client\'
import { create } from \'zustand\'
import { persist } from \'zustand/middleware\'
import type { Thread, Message, Source, RetrievedImage, FocusMode, CitationEntry, DocumentInfo } from \'@/lib/types\'
import { generateId } from \'@/lib/utils\'

interface SearchState {
  // Thread
  currentThreadId: string | null
  threads: Thread[]
  isLoading: boolean
  isStreaming: boolean
  streamText: string
  abortController: AbortController | null

  // Settings
  model: string
  focus: FocusMode
  sidebarOpen: boolean
  useHyde: boolean
  useDualPath: boolean
  backendOnline: boolean

  // Documents (session state)
  documents: DocumentInfo[]
  selectedDocIds: string[]  // for scoped search

  // Live stream state
  _currentSources:  Source[]
  _currentImages:   RetrievedImage[]
  _currentCitations: Record<string, CitationEntry>

  // Actions
  setModel:          (m: string) => void
  setFocus:          (f: FocusMode) => void
  setUseHyde:        (v: boolean) => void
  setUseDualPath:    (v: boolean) => void
  toggleSidebar:     () => void
  setSidebarOpen:    (v: boolean) => void
  setBackendOnline:  (v: boolean) => void
  setDocuments:      (docs: DocumentInfo[]) => void
  toggleDocSelected: (id: string) => void
  selectAllDocs:     () => void
  clearDocSelection: () => void

  startStream:    () => AbortController
  appendStream:   (delta: string) => void
  endStream:      (sources: Source[], images: RetrievedImage[], related: string[], citations: Record<string, CitationEntry>, meta?: Record<string, unknown>) => void
  cancelStream:   () => void

  setSources:    (s: Source[]) => void
  setImages:     (i: RetrievedImage[]) => void
  setCitations:  (c: Record<string, CitationEntry>) => void

  createThread:     (query: string) => Thread
  addUserMessage:   (threadId: string, content: string) => void
  getCurrentThread: () => Thread | null
  deleteThread:     (id: string) => void
  clearAll:         () => void
}

export const useSearchStore = create<SearchState>()(
  persist(
    (set, get) => ({
      currentThreadId: null,
      threads: [],
      isLoading: false,
      isStreaming: false,
      streamText: \'\',
      abortController: null,
      model: \'openrouter/free\',
      focus: \'all\',
      sidebarOpen: true,
      useHyde: true,
      useDualPath: true,
      backendOnline: false,
      documents: [],
      selectedDocIds: [],
      _currentSources: [],
      _currentImages: [],
      _currentCitations: {},

      setModel:         (model) => set({ model }),
      setFocus:         (focus) => set({ focus }),
      setUseHyde:       (v) => set({ useHyde: v }),
      setUseDualPath:   (v) => set({ useDualPath: v }),
      toggleSidebar:    () => set((s) => ({ sidebarOpen: !s.sidebarOpen })),
      setSidebarOpen:   (v) => set({ sidebarOpen: v }),
      setBackendOnline: (v) => set({ backendOnline: v }),
      setDocuments:     (docs) => set({ documents: docs }),
      toggleDocSelected: (id) => set((s) => ({
        selectedDocIds: s.selectedDocIds.includes(id)
          ? s.selectedDocIds.filter((d) => d !== id)
          : [...s.selectedDocIds, id],
      })),
      selectAllDocs:  () => set((s) => ({ selectedDocIds: s.documents.map((d) => d.doc_id) })),
      clearDocSelection: () => set({ selectedDocIds: [] }),
      setSources:    (s) => set({ _currentSources: s }),
      setImages:     (i) => set({ _currentImages: i }),
      setCitations:  (c) => set({ _currentCitations: c }),

      startStream: () => {
        const ac = new AbortController()
        set({ isStreaming: true, isLoading: false, streamText: \'\',
              abortController: ac, _currentSources: [], _currentImages: [],
              _currentCitations: {} })
        return ac
      },

      appendStream: (delta) => set((s) => ({ streamText: s.streamText + delta })),

      endStream: (sources, images, related, citations, meta) => {
        const { streamText, currentThreadId, threads } = get()
        if (!currentThreadId) return
        const msg: Message = {
          id: generateId(), role: \'assistant\',
          content: streamText, sources, images,
          related_questions: related,
          citation_map: citations,
          timestamp: new Date(),
          meta: meta as Message[\'meta\'],
        }
        const updated = threads.map((t) =>
          t.id === currentThreadId
            ? { ...t, messages: [...t.messages, msg], updatedAt: new Date() }
            : t
        )
        set({ isStreaming: false, streamText: \'\', threads: updated,
              abortController: null, _currentSources: [], _currentImages: [],
              _currentCitations: {} })
      },

      cancelStream: () => {
        get().abortController?.abort()
        set({ isStreaming: false, isLoading: false, streamText: \'\', abortController: null })
      },

      createThread: (query) => {
        const thread: Thread = {
          id: generateId(),
          title: query.length > 60 ? query.slice(0, 60) + \'…\' : query,
          messages: [], createdAt: new Date(), updatedAt: new Date(),
          model: get().model, focus: get().focus,
        }
        set((s) => ({ threads: [thread, ...s.threads], currentThreadId: thread.id }))
        return thread
      },

      addUserMessage: (threadId, content) => {
        const msg: Message = { id: generateId(), role: \'user\', content, timestamp: new Date() }
        set((s) => ({
          threads: s.threads.map((t) =>
            t.id === threadId
              ? { ...t, messages: [...t.messages, msg], updatedAt: new Date() }
              : t
          ),
          isLoading: true,
        }))
      },

      getCurrentThread: () => {
        const { threads, currentThreadId } = get()
        return threads.find((t) => t.id === currentThreadId) ?? null
      },

      deleteThread: (id) => set((s) => ({
        threads: s.threads.filter((t) => t.id !== id),
        currentThreadId: s.currentThreadId === id ? null : s.currentThreadId,
      })),

      clearAll: () => set({ threads: [], currentThreadId: null }),
    }),
    {
      name: \'ragraph-store\',
      partialize: (s) => ({
        threads: s.threads,
        model: s.model,
        focus: s.focus,
        sidebarOpen: s.sidebarOpen,
        useHyde: s.useHyde,
        useDualPath: s.useDualPath,
      }),
    }
  )
)
''', base=FRONTEND)

# ═══════════════════════════════════════════════════════════
# FRONTEND: src/components/DocumentManager.tsx
# ═══════════════════════════════════════════════════════════
w("components/DocumentManager.tsx", '''
\'use client\'
import { useState, useRef, useEffect, useCallback } from \'react\'
import { motion, AnimatePresence } from \'framer-motion\'
import { Upload, Trash2, Eye, FileText, Loader2, CheckCircle, AlertCircle, X } from \'lucide-react\'
import { useSearchStore } from \'@/store/useSearchStore\'
import { uploadDocument, pollIngestionStatus, listDocuments, deleteDocument, getDocumentContent } from \'@/lib/api\'
import type { DocumentInfo } from \'@/lib/types\'

const STATUS_LABELS: Record<string, string> = {
  queued:    \'Queued\',
  parsing:   \'Parsing…\',
  embedding: \'Embedding…\',
  indexing:  \'Indexing…\',
  done:      \'Ready\',
  error:     \'Error\',
}

const STATUS_COLOR: Record<string, string> = {
  queued:    \'var(--text-muted)\',
  parsing:   \'var(--accent-amber)\',
  embedding: \'var(--accent-blue)\',
  indexing:  \'var(--accent-blue)\',
  done:      \'var(--accent-green)\',
  error:     \'#ef4444\',
}

interface Props {
  onClose?: () => void
}

export function DocumentManager({ onClose }: Props) {
  const { documents, setDocuments, selectedDocIds, toggleDocSelected } = useSearchStore()
  const [uploading, setUploading] = useState(false)
  const [uploadPct, setUploadPct] = useState(0)
  const [viewer, setViewer] = useState<{ doc: DocumentInfo; content: string } | null>(null)
  const [loadingView, setLoadingView] = useState<string | null>(null)
  const fileRef = useRef<HTMLInputElement>(null)

  // Refresh document list
  const refresh = useCallback(async () => {
    try {
      const data = await listDocuments()
      setDocuments(data.documents)
    } catch { /* backend might not be ready */ }
  }, [setDocuments])

  useEffect(() => {
    refresh()
    const interval = setInterval(refresh, 3000)
    return () => clearInterval(interval)
  }, [refresh])

  const handleUpload = async (e: React.ChangeEvent<HTMLInputElement>) => {
    const files = Array.from(e.target.files ?? [])
    if (!files.length) return
    e.target.value = \'\'
    setUploading(true)

    for (const file of files) {
      try {
        setUploadPct(0)
        const res = await uploadDocument(file, setUploadPct)
        // Poll until done
        let attempts = 0
        while (attempts < 120) {
          await new Promise(r => setTimeout(r, 1500))
          const status = await pollIngestionStatus(res.doc_id)
          await refresh()
          if (status.status === \'done\' || status.status === \'error\') break
          attempts++
        }
      } catch (err) {
        console.error(\'Upload failed:\', err)
      }
    }
    setUploading(false)
    await refresh()
  }

  const handleDelete = async (doc: DocumentInfo) => {
    if (!confirm(`Delete [Doc ${doc.number}] ${doc.filename}?`)) return
    try {
      await deleteDocument(doc.doc_id)
      await refresh()
    } catch (err) {
      alert(\'Delete failed: \' + err)
    }
  }

  const handleView = async (doc: DocumentInfo) => {
    setLoadingView(doc.doc_id)
    try {
      const content = await getDocumentContent(doc.doc_id)
      setViewer({ doc, content })
    } catch (err) {
      alert(\'Could not load document: \' + err)
    } finally {
      setLoadingView(null)
    }
  }

  const formatSize = (bytes: number) => {
    if (bytes < 1024) return `${bytes}B`
    if (bytes < 1024 * 1024) return `${(bytes / 1024).toFixed(1)}KB`
    return `${(bytes / 1024 / 1024).toFixed(1)}MB`
  }

  return (
    <div style={{ display: \'flex\', flexDirection: \'column\', height: \'100%\' }}>
      {/* Header */}
      <div style={{
        padding: \'16px 16px 12px\',
        borderBottom: \'1px solid var(--border)\',
        display: \'flex\', alignItems: \'center\', gap: 10,
      }}>
        <h2 style={{ flex: 1, fontSize: 15, fontWeight: 500 }}>Documents</h2>
        <span style={{ fontSize: 11, color: \'var(--text-muted)\' }}>
          {documents.length} file{documents.length !== 1 ? \'s\' : \'\'}
        </span>
        {onClose && (
          <button onClick={onClose} className="btn-ghost" style={{ padding: \'4px 6px\' }}>
            <X size={14} />
          </button>
        )}
      </div>

      {/* Upload button */}
      <div style={{ padding: \'12px 12px 6px\' }}>
        <input
          ref={fileRef} type="file" multiple
          accept=".pdf,.txt,.md,.docx,.pptx"
          style={{ display: \'none\' }}
          onChange={handleUpload}
        />
        <button
          onClick={() => fileRef.current?.click()}
          disabled={uploading}
          style={{
            width: \'100%\', display: \'flex\', alignItems: \'center\',
            justifyContent: \'center\', gap: 8,
            padding: \'9px 12px\',
            background: \'var(--accent-blue)\', color: \'white\',
            border: \'none\', borderRadius: 10, cursor: uploading ? \'not-allowed\' : \'pointer\',
            fontSize: 13, fontWeight: 500, fontFamily: \'var(--font-body)\',
            opacity: uploading ? 0.7 : 1,
          }}
        >
          {uploading
            ? <><Loader2 size={13} style={{ animation: \'spin 1s linear infinite\' }} />
                Processing {uploadPct}%…</>
            : <><Upload size={13} /> Upload documents</>
          }
        </button>
      </div>

      {/* Document list */}
      <div style={{ flex: 1, overflow: \'auto\', padding: \'4px 8px\' }}>
        {documents.length === 0 ? (
          <div style={{
            textAlign: \'center\', padding: \'40px 16px\',
            color: \'var(--text-muted)\', fontSize: 12.5, lineHeight: 1.6,
          }}>
            <FileText size={28} style={{ margin: \'0 auto 12px\', opacity: 0.4, display: \'block\' }} />
            No documents yet.<br />Upload PDFs, TXT, DOCX, PPTX.
          </div>
        ) : (
          <AnimatePresence>
            {documents.map((doc) => {
              const isSelected = selectedDocIds.includes(doc.doc_id)
              const isProcessing = [\'queued\', \'parsing\', \'embedding\', \'indexing\'].includes(doc.status)
              return (
                <motion.div
                  key={doc.doc_id}
                  initial={{ opacity: 0, x: -10 }}
                  animate={{ opacity: 1, x: 0 }}
                  exit={{ opacity: 0, x: 10 }}
                  style={{
                    display: \'flex\', alignItems: \'flex-start\', gap: 8,
                    padding: \'8px 10px\', borderRadius: 10, marginBottom: 4,
                    background: isSelected ? \'var(--accent-blue-light)\' : \'transparent\',
                    border: isSelected ? \'1px solid var(--border-focus)\' : \'1px solid transparent\',
                    cursor: \'pointer\', transition: \'background 0.15s, border 0.15s\',
                  }}
                  onClick={() => doc.status === \'done\' && toggleDocSelected(doc.doc_id)}
                >
                  {/* Number badge */}
                  <div style={{
                    flexShrink: 0, width: 22, height: 22, borderRadius: 6,
                    background: doc.status === \'done\' ? \'var(--accent-blue)\' : \'var(--bg-hover)\',
                    color: doc.status === \'done\' ? \'white\' : \'var(--text-muted)\',
                    display: \'flex\', alignItems: \'center\', justifyContent: \'center\',
                    fontSize: 11, fontWeight: 700, marginTop: 1,
                  }}>
                    {doc.number}
                  </div>

                  {/* Info */}
                  <div style={{ flex: 1, minWidth: 0 }}>
                    <p style={{
                      fontSize: 12.5, fontWeight: 500, color: \'var(--text-primary)\',
                      overflow: \'hidden\', textOverflow: \'ellipsis\', whiteSpace: \'nowrap\',
                      lineHeight: 1.4,
                    }}>
                      {doc.filename}
                    </p>
                    <div style={{ display: \'flex\', alignItems: \'center\', gap: 6, marginTop: 3 }}>
                      {isProcessing
                        ? <Loader2 size={10} style={{
                            color: STATUS_COLOR[doc.status],
                            animation: \'spin 1s linear infinite\',
                          }} />
                        : doc.status === \'done\'
                          ? <CheckCircle size={10} style={{ color: STATUS_COLOR.done }} />
                          : <AlertCircle size={10} style={{ color: STATUS_COLOR.error }} />
                      }
                      <span style={{ fontSize: 10.5, color: STATUS_COLOR[doc.status] }}>
                        {STATUS_LABELS[doc.status]}
                      </span>
                      {doc.status === \'done\' && (
                        <>
                          <span style={{ fontSize: 10, color: \'var(--text-muted)\' }}>·</span>
                          <span style={{ fontSize: 10, color: \'var(--text-muted)\' }}>
                            {doc.node_count} chunks · {formatSize(doc.size_bytes)}
                          </span>
                        </>
                      )}
                    </div>
                  </div>

                  {/* Actions */}
                  <div style={{ display: \'flex\', gap: 4, flexShrink: 0 }}>
                    {doc.status === \'done\' && (
                      <button
                        onClick={(e) => { e.stopPropagation(); handleView(doc) }}
                        style={{
                          background: \'none\', border: \'none\', cursor: \'pointer\',
                          color: \'var(--text-muted)\', padding: 3, borderRadius: 4,
                          display: \'flex\',
                        }}
                        title="View document"
                      >
                        {loadingView === doc.doc_id
                          ? <Loader2 size={12} style={{ animation: \'spin 1s linear infinite\' }} />
                          : <Eye size={12} />
                        }
                      </button>
                    )}
                    <button
                      onClick={(e) => { e.stopPropagation(); handleDelete(doc) }}
                      style={{
                        background: \'none\', border: \'none\', cursor: \'pointer\',
                        color: \'var(--text-muted)\', padding: 3, borderRadius: 4,
                        display: \'flex\',
                      }}
                      title="Delete document"
                    >
                      <Trash2 size={12} />
                    </button>
                  </div>
                </motion.div>
              )
            })}
          </AnimatePresence>
        )}
      </div>

      {/* Selection indicator */}
      {documents.filter(d => d.status === \'done\').length > 0 && (
        <div style={{
          padding: \'8px 12px\', borderTop: \'1px solid var(--border)\',
          fontSize: 11, color: \'var(--text-muted)\',
          display: \'flex\', alignItems: \'center\', gap: 8,
        }}>
          <span>
            {selectedDocIds.length === 0
              ? \'Click docs to scope search\'
              : `${selectedDocIds.length} selected for search`}
          </span>
          {selectedDocIds.length > 0 && (
            <button
              onClick={() => useSearchStore.getState().clearDocSelection()}
              style={{ fontSize: 10, color: \'var(--accent-blue)\', background: \'none\',
                       border: \'none\', cursor: \'pointer\' }}
            >
              clear
            </button>
          )}
        </div>
      )}

      {/* Document viewer modal */}
      <AnimatePresence>
        {viewer && (
          <motion.div
            initial={{ opacity: 0 }}
            animate={{ opacity: 1 }}
            exit={{ opacity: 0 }}
            style={{
              position: \'fixed\', inset: 0, background: \'rgba(0,0,0,0.7)\',
              zIndex: 200, display: \'flex\', alignItems: \'center\',
              justifyContent: \'center\', padding: 24,
            }}
            onClick={() => setViewer(null)}
          >
            <motion.div
              initial={{ scale: 0.95 }}
              animate={{ scale: 1 }}
              exit={{ scale: 0.95 }}
              onClick={(e) => e.stopPropagation()}
              style={{
                background: \'var(--bg-card)\', borderRadius: 16,
                width: \'100%\', maxWidth: 760, maxHeight: \'80vh\',
                display: \'flex\', flexDirection: \'column\',
                boxShadow: \'0 25px 80px rgba(0,0,0,0.5)\',
              }}
            >
              {/* Viewer header */}
              <div style={{
                padding: \'14px 18px\', borderBottom: \'1px solid var(--border)\',
                display: \'flex\', alignItems: \'center\', gap: 10,
              }}>
                <span style={{
                  background: \'var(--accent-blue)\', color: \'white\',
                  borderRadius: 6, padding: \'2px 8px\', fontSize: 12, fontWeight: 700,
                }}>
                  Doc {viewer.doc.number}
                </span>
                <span style={{ flex: 1, fontSize: 14, fontWeight: 500 }}>
                  {viewer.doc.filename}
                </span>
                <button onClick={() => setViewer(null)}
                  style={{ background: \'none\', border: \'none\', cursor: \'pointer\',
                           color: \'var(--text-muted)\', display: \'flex\' }}>
                  <X size={16} />
                </button>
              </div>
              {/* Content */}
              <pre style={{
                flex: 1, overflow: \'auto\', padding: \'16px 18px\',
                fontFamily: \'var(--font-mono)\', fontSize: 12, lineHeight: 1.7,
                color: \'var(--text-secondary)\', whiteSpace: \'pre-wrap\', wordBreak: \'break-word\',
              }}>
                {viewer.content}
              </pre>
            </motion.div>
          </motion.div>
        )}
      </AnimatePresence>

      <style>{`@keyframes spin { to { transform: rotate(360deg); } }`}</style>
    </div>
  )
}
''', base=FRONTEND)

# ═══════════════════════════════════════════════════════════
# FRONTEND: src/components/CitationMap.tsx
# ═══════════════════════════════════════════════════════════
w("components/CitationMap.tsx", '''
\'use client\'
import type { CitationEntry } from \'@/lib/types\'

interface Props {
  citationMap: Record<string, CitationEntry>
}

export function CitationMap({ citationMap }: Props) {
  const entries = Object.entries(citationMap)
  if (!entries.length) return null

  // Deduplicate by doc_number
  const seen = new Set<number>()
  const unique = entries.filter(([, v]) => {
    if (seen.has(v.doc_number)) return false
    seen.add(v.doc_number)
    return true
  })

  return (
    <div style={{
      marginTop: 12, padding: \'10px 14px\',
      background: \'var(--bg-secondary)\',
      border: \'1px solid var(--border)\',
      borderRadius: 10, fontSize: 12,
    }}>
      <p style={{
        fontSize: 10.5, fontWeight: 600, letterSpacing: \'0.06em\',
        textTransform: \'uppercase\', color: \'var(--text-muted)\', marginBottom: 8,
      }}>
        Sources used
      </p>
      <div style={{ display: \'flex\', flexWrap: \'wrap\', gap: 6 }}>
        {unique.map(([, entry]) => (
          <div key={entry.doc_number} style={{
            display: \'inline-flex\', alignItems: \'center\', gap: 6,
            background: \'var(--bg-card)\',
            border: \'1px solid var(--border)\',
            borderRadius: 8, padding: \'4px 10px\',
          }}>
            <span style={{
              background: \'var(--accent-blue)\', color: \'white\',
              borderRadius: 4, padding: \'1px 6px\',
              fontSize: 10, fontWeight: 700,
            }}>
              Doc {entry.doc_number}
            </span>
            <span style={{ color: \'var(--text-secondary)\', fontSize: 12 }}>
              {entry.filename.length > 30
                ? entry.filename.slice(0, 28) + \'…\'
                : entry.filename}
            </span>
          </div>
        ))}
      </div>
    </div>
  )
}
''', base=FRONTEND)

# ═══════════════════════════════════════════════════════════
# FRONTEND: update AnswerCard to show CitationMap
# ═══════════════════════════════════════════════════════════
answercard_path = os.path.join(FRONTEND, "components", "AnswerCard.tsx")
if os.path.exists(answercard_path):
    with open(answercard_path, "r") as f:
        ac = f.read()
    # Add CitationMap import + prop
    if "CitationMap" not in ac:
        ac = ac.replace(
            "import { SourceCard } from './SourceCard'",
            "import { SourceCard } from './SourceCard'\nimport { CitationMap } from './CitationMap'"
        )
        ac = ac.replace(
            "  onRetry?: () => void\n}",
            "  citationMap?: Record<string, import('@/lib/types').CitationEntry>\n  onRetry?: () => void\n}"
        )
        ac = ac.replace(
            "  isStreaming, meta, onFollowUp, onRetry,\n}: Props)",
            "  isStreaming, meta, citationMap, onFollowUp, onRetry,\n}: Props)"
        )
        # Add CitationMap after related questions block
        ac = ac.replace(
            "    </motion.div>\n  )\n}",
            "      {citationMap && Object.keys(citationMap).length > 0 && (\n        <CitationMap citationMap={citationMap} />\n      )}\n    </motion.div>\n  )\n}"
        )
        with open(answercard_path, "w") as f:
            f.write(ac)
        print("  AnswerCard.tsx: CitationMap added")

# ═══════════════════════════════════════════════════════════
# FRONTEND: update Sidebar to include DocumentManager
# ═══════════════════════════════════════════════════════════
w("components/Sidebar.tsx", '''
\'use client\'
import { motion, AnimatePresence } from \'framer-motion\'
import { Plus, Trash2, MessageSquare, X, Files } from \'lucide-react\'
import { useSearchStore } from \'@/store/useSearchStore\'
import { DocumentManager } from \'./DocumentManager\'
import { formatRelativeTime } from \'@/lib/utils\'
import { useState } from \'react\'

interface Props { onNewSearch: () => void }

export function Sidebar({ onNewSearch }: Props) {
  const { threads, currentThreadId, sidebarOpen, toggleSidebar, deleteThread } = useSearchStore()
  const store = useSearchStore()
  const [tab, setTab] = useState<\'threads\' | \'docs\'>(\'docs\')

  return (
    <AnimatePresence>
      {sidebarOpen && (
        <motion.aside
          initial={{ width: 0, opacity: 0 }}
          animate={{ width: 270, opacity: 1 }}
          exit={{ width: 0, opacity: 0 }}
          transition={{ type: \'spring\', stiffness: 300, damping: 30 }}
          style={{
            height: \'100vh\', background: \'var(--bg-secondary)\',
            borderRight: \'1px solid var(--border)\',
            display: \'flex\', flexDirection: \'column\',
            overflow: \'hidden\', flexShrink: 0,
            position: \'sticky\', top: 0,
          }}
        >
          {/* Header */}
          <div style={{
            padding: \'16px 16px 12px\',
            display: \'flex\', alignItems: \'center\', gap: 10,
            borderBottom: \'1px solid var(--border)\',
          }}>
            <div style={{ flex: 1 }}>
              <h1 style={{
                fontFamily: \'var(--font-display)\', fontSize: 18,
                fontWeight: 400, letterSpacing: \'-0.02em\',
              }}>RAGraph</h1>
              <p style={{ fontSize: 10.5, color: \'var(--text-muted)\', marginTop: 1 }}>
                Hierarchical RAG
              </p>
            </div>
            <button onClick={toggleSidebar} className="btn-ghost" style={{ padding: 6, borderRadius: 8 }}>
              <X size={14} />
            </button>
          </div>

          {/* Tabs */}
          <div style={{
            display: \'flex\', borderBottom: \'1px solid var(--border)\',
            padding: \'8px 12px 0\', gap: 4,
          }}>
            {[
              { key: \'docs\',    label: \'Documents\', icon: <Files size={12} /> },
              { key: \'threads\', label: \'History\',   icon: <MessageSquare size={12} /> },
            ].map(({ key, label, icon }) => (
              <button
                key={key}
                onClick={() => setTab(key as any)}
                style={{
                  display: \'flex\', alignItems: \'center\', gap: 5,
                  padding: \'6px 12px\', borderRadius: \'8px 8px 0 0\',
                  background: tab === key ? \'var(--bg-card)\' : \'transparent\',
                  border: tab === key ? \'1px solid var(--border)\' : \'1px solid transparent\',
                  borderBottom: tab === key ? \'1px solid var(--bg-card)\' : \'1px solid transparent\',
                  cursor: \'pointer\', fontSize: 12, fontWeight: tab === key ? 500 : 400,
                  color: tab === key ? \'var(--text-primary)\' : \'var(--text-muted)\',
                  fontFamily: \'var(--font-body)\',
                  marginBottom: -1,
                }}
              >
                {icon}{label}
              </button>
            ))}
          </div>

          {/* Tab content */}
          <div style={{ flex: 1, overflow: \'hidden\', display: \'flex\', flexDirection: \'column\' }}>
            {tab === \'docs\' ? (
              <DocumentManager />
            ) : (
              <>
                <div style={{ padding: \'10px 12px 6px\' }}>
                  <button onClick={onNewSearch} style={{
                    width: \'100%\', display: \'flex\', alignItems: \'center\', gap: 8,
                    padding: \'9px 12px\', background: \'var(--accent-blue)\', color: \'white\',
                    border: \'none\', borderRadius: 10, cursor: \'pointer\',
                    fontSize: 13, fontWeight: 500, fontFamily: \'var(--font-body)\',
                  }}>
                    <Plus size={14} /> New search
                  </button>
                </div>
                <div style={{ flex: 1, overflow: \'auto\', padding: \'4px 8px\' }}>
                  {threads.length === 0 ? (
                    <p style={{ fontSize: 12, color: \'var(--text-muted)\', textAlign: \'center\',
                                marginTop: 32, lineHeight: 1.6 }}>
                      Search history<br />appears here
                    </p>
                  ) : (
                    threads.map((thread) => (
                      <motion.div key={thread.id}
                        initial={{ opacity: 0, x: -10 }} animate={{ opacity: 1, x: 0 }}
                        style={{
                          display: \'flex\', alignItems: \'flex-start\', gap: 8,
                          padding: \'8px 10px\', borderRadius: 10, cursor: \'pointer\',
                          background: thread.id === currentThreadId ? \'var(--bg-hover)\' : \'transparent\',
                          marginBottom: 2,
                        }}
                        onClick={() => store.setState({ currentThreadId: thread.id })}
                        whileHover={{ background: \'var(--bg-hover)\' } as any}
                      >
                        <MessageSquare size={13} style={{ color: \'var(--text-muted)\', flexShrink: 0, marginTop: 2 }} />
                        <div style={{ flex: 1, minWidth: 0 }}>
                          <p style={{
                            fontSize: 12.5, fontWeight: thread.id === currentThreadId ? 500 : 400,
                            overflow: \'hidden\', textOverflow: \'ellipsis\', whiteSpace: \'nowrap\',
                          }}>
                            {thread.title}
                          </p>
                          <p style={{ fontSize: 10.5, color: \'var(--text-muted)\', marginTop: 2 }}>
                            {formatRelativeTime(new Date(thread.updatedAt))}
                          </p>
                        </div>
                        <button
                          onClick={(e) => { e.stopPropagation(); deleteThread(thread.id) }}
                          style={{ background: \'none\', border: \'none\', cursor: \'pointer\',
                                   color: \'var(--text-muted)\', display: \'flex\', opacity: 0 }}
                          className="delete-btn"
                        >
                          <Trash2 size={12} />
                        </button>
                      </motion.div>
                    ))
                  )}
                </div>
              </>
            )}
          </div>
        </motion.aside>
      )}
      <style>{`div:hover .delete-btn { opacity: 1 !important; }`}</style>
    </AnimatePresence>
  )
}
''', base=FRONTEND)

# ═══════════════════════════════════════════════════════════
# FRONTEND: update page.tsx to wire everything
# ═══════════════════════════════════════════════════════════
w("app/page.tsx", '''
\'use client\'
import { useCallback, useState, useEffect, useRef } from \'react\'
import { motion, AnimatePresence } from \'framer-motion\'
import { SearchBar }         from \'@/components/SearchBar\'
import { Sidebar }           from \'@/components/Sidebar\'
import { Navbar }            from \'@/components/Navbar\'
import { AnswerCard }        from \'@/components/AnswerCard\'
import { ThinkingIndicator } from \'@/components/ThinkingIndicator\'
import { useSearchStore }    from \'@/store/useSearchStore\'
import { streamSearch, healthCheck } from \'@/lib/api\'
import type { Source, RetrievedImage, SearchRequest, CitationEntry } from \'@/lib/types\'

const SUGGESTIONS = [
  \'What are the main topics in my documents?\',
  \'Summarize the key findings\',
  \'What methodology was used?\',
  \'What figures are referenced?\',
]

export default function Home() {
  const store = useSearchStore()
  const {
    currentThreadId, threads, isLoading, isStreaming, streamText,
    model, focus, useHyde, useDualPath,
    _currentSources, _currentImages, _currentCitations,
    selectedDocIds, backendOnline,
    startStream, appendStream, endStream,
    addUserMessage, createThread,
    setBackendOnline, setSources, setImages, setCitations,
  } = store

  const [related,     setRelated]     = useState<string[]>([])
  const [citations,   setCitationsLocal] = useState<Record<string, CitationEntry>>({})
  const [thinkStep,   setThinkStep]   = useState(0)
  const bottomRef = useRef<HTMLDivElement>(null)

  const currentThread = threads.find((t) => t.id === currentThreadId) ?? null
  const isHomePage    = !currentThreadId

  useEffect(() => {
    healthCheck().then(() => setBackendOnline(true)).catch(() => setBackendOnline(false))
  }, [setBackendOnline])

  useEffect(() => {
    if (isStreaming) bottomRef.current?.scrollIntoView({ behavior: \'smooth\' })
  }, [streamText, isStreaming])

  useEffect(() => {
    if (!isLoading) { setThinkStep(0); return }
    const timers = [400, 900, 1600].map((d, i) => setTimeout(() => setThinkStep(i + 1), d))
    return () => timers.forEach(clearTimeout)
  }, [isLoading])

  const handleSearch = useCallback(async (query: string, imageBase64?: string) => {
    let threadId = currentThreadId
    if (!threadId || isHomePage) {
      const thread = createThread(query)
      threadId = thread.id
    }
    addUserMessage(threadId!, query)
    setRelated([])
    setCitationsLocal({})

    const ac = startStream()
    const request: SearchRequest = {
      query, model, focus,
      thread_id: threadId ?? undefined,
      image: imageBase64,
      use_hyde: useHyde, use_dual_path: useDualPath,
      conversation_history: currentThread?.messages.map((m) => ({
        role: m.role, content: m.content,
      })),
    }

    let finalSources:   Source[]                         = []
    let finalImages:    RetrievedImage[]                 = []
    let finalRelated:   string[]                         = []
    let finalCitations: Record<string, CitationEntry>    = {}
    let finalMeta:      Record<string, unknown>          = {}

    await streamSearch(request, {
      onText:      appendStream,
      onSources:   (s) => { finalSources   = s; setSources(s) },
      onImages:    (i) => { finalImages    = i; setImages(i)  },
      onCitations: (c) => { finalCitations = c; setCitations(c); setCitationsLocal(c) },
      onRelated:   (r) => { finalRelated   = r; setRelated(r) },
      onDone:      (m) => { finalMeta      = m },
      onError:     (e) => { console.error(\'Stream error:\', e); endStream([], [], [], {}, {}) },
    }, ac.signal)

    endStream(finalSources, finalImages, finalRelated, finalCitations, finalMeta)
  }, [currentThreadId, isHomePage, currentThread, model, focus, useHyde, useDualPath,
      createThread, addUserMessage, startStream, appendStream, endStream,
      setSources, setImages, setCitations])

  const handleNewSearch = useCallback(() => {
    store.setState({ currentThreadId: null })
  }, [store])

  const messages = currentThread?.messages ?? []

  return (
    <div style={{ display: \'flex\', height: \'100vh\', overflow: \'hidden\', background: \'var(--bg-primary)\' }}>
      <Sidebar onNewSearch={handleNewSearch} />

      <div style={{ flex: 1, display: \'flex\', flexDirection: \'column\', overflow: \'hidden\' }}>
        <Navbar />

        {!backendOnline && (
          <motion.div initial={{ height: 0 }} animate={{ height: \'auto\' }} style={{
            background: \'var(--bg-secondary)\', borderBottom: \'1px solid var(--border)\',
            padding: \'8px 24px\', fontSize: 12.5, color: \'var(--text-secondary)\',
            display: \'flex\', alignItems: \'center\', gap: 8,
          }}>
            <span style={{ color: \'#f59e0b\' }}>⚠</span>
            Backend offline — run{` `}
            <code style={{ background: \'var(--bg-hover)\', padding: \'1px 6px\', borderRadius: 4 }}>
              uvicorn app.main:app --reload --port 8000
            </code>
          </motion.div>
        )}

        {/* Selected docs indicator */}
        {selectedDocIds.length > 0 && (
          <div style={{
            padding: \'6px 24px\', background: \'var(--accent-blue-light)\',
            borderBottom: \'1px solid var(--border-focus)\',
            fontSize: 12, color: \'var(--accent-blue)\',
            display: \'flex\', alignItems: \'center\', gap: 8,
          }}>
            <span>🔍 Searching in {selectedDocIds.length} selected document{selectedDocIds.length > 1 ? \'s\' : \'\'}</span>
            <button onClick={() => store.clearDocSelection()}
              style={{ fontSize: 11, color: \'var(--accent-blue)\', background: \'none\',
                       border: \'none\', cursor: \'pointer\', textDecoration: \'underline\' }}>
              search all
            </button>
          </div>
        )}

        <main style={{ flex: 1, overflow: \'auto\', padding: \'0 0 24px\' }}>
          <AnimatePresence mode="wait">
            {isHomePage ? (
              <motion.div key="home" initial={{ opacity: 0 }} animate={{ opacity: 1 }}
                style={{ maxWidth: 680, margin: \'0 auto\', padding: \'80px 24px 120px\',
                         display: \'flex\', flexDirection: \'column\', alignItems: \'center\', gap: 32 }}>
                <div style={{ textAlign: \'center\' }}>
                  <motion.h1 initial={{ opacity: 0, y: 20 }} animate={{ opacity: 1, y: 0 }}
                    transition={{ delay: 0.1 }}
                    style={{ fontFamily: \'var(--font-display)\', fontSize: \'clamp(36px,5vw,52px)\',
                             fontWeight: 400, letterSpacing: \'-0.03em\', lineHeight: 1.15, marginBottom: 14 }}>
                    Ask anything,{` `}
                    <span className="gradient-text">find everything.</span>
                  </motion.h1>
                  <motion.p initial={{ opacity: 0, y: 16 }} animate={{ opacity: 1, y: 0 }}
                    transition={{ delay: 0.2 }}
                    style={{ fontSize: 16, color: \'var(--text-secondary)\', maxWidth: 440,
                             margin: \'0 auto\', lineHeight: 1.6 }}>
                    Upload documents in the sidebar, then ask questions.
                    Responses cite <strong>[Doc N]</strong> so you always know the source.
                  </motion.p>
                </div>

                <motion.div initial={{ opacity: 0, y: 16 }} animate={{ opacity: 1, y: 0 }}
                  transition={{ delay: 0.25 }} style={{ width: \'100%\' }}>
                  <SearchBar onSearch={handleSearch} />
                </motion.div>

                <motion.div initial={{ opacity: 0 }} animate={{ opacity: 1 }}
                  transition={{ delay: 0.35 }} style={{ width: \'100%\' }}>
                  <p className="section-label" style={{ marginBottom: 12, textAlign: \'center\' }}>
                    Try asking
                  </p>
                  <div style={{ display: \'grid\', gridTemplateColumns: \'repeat(auto-fit,minmax(240px,1fr))\', gap: 10 }}>
                    {SUGGESTIONS.map((s, i) => (
                      <motion.button key={i}
                        initial={{ opacity: 0, y: 8 }} animate={{ opacity: 1, y: 0 }}
                        transition={{ delay: 0.4 + i * 0.07 }}
                        onClick={() => handleSearch(s)}
                        style={{ padding: \'12px 16px\', background: \'var(--bg-card)\',
                                 border: \'1px solid var(--border)\', borderRadius: 12,
                                 cursor: \'pointer\', fontSize: 13.5, color: \'var(--text-primary)\',
                                 textAlign: \'left\', fontFamily: \'var(--font-body)\', lineHeight: 1.45 }}
                        whileHover={{ y: -2, boxShadow: \'var(--shadow-md)\' }}
                      >
                        {s}
                      </motion.button>
                    ))}
                  </div>
                </motion.div>
              </motion.div>
            ) : (
              <motion.div key={currentThreadId} initial={{ opacity: 0 }} animate={{ opacity: 1 }}
                style={{ maxWidth: 760, margin: \'0 auto\', padding: \'24px 24px 0\' }}>
                {messages.map((msg) => (
                  <div key={msg.id} style={{ marginBottom: 28 }}>
                    {msg.role === \'user\' ? (
                      <motion.div initial={{ opacity: 0, y: 8 }} animate={{ opacity: 1, y: 0 }}
                        style={{ display: \'flex\', justifyContent: \'flex-end\', marginBottom: 16 }}>
                        <div style={{
                          background: \'var(--bg-secondary)\', border: \'1px solid var(--border)\',
                          borderRadius: \'16px 16px 4px 16px\', padding: \'12px 18px\',
                          maxWidth: \'80%\', fontSize: 15, lineHeight: 1.55,
                        }}>
                          {msg.content}
                        </div>
                      </motion.div>
                    ) : (
                      <AnswerCard
                        content={msg.content}
                        sources={msg.sources ?? []}
                        images={msg.images ?? []}
                        relatedQuestions={msg.related_questions ?? []}
                        citationMap={msg.citation_map}
                        meta={msg.meta}
                        onFollowUp={handleSearch}
                      />
                    )}
                  </div>
                ))}

                {(isLoading || isStreaming) && (
                  <div style={{ marginBottom: 28 }}>
                    <ThinkingIndicator visible={isLoading} step={thinkStep} />
                    {isStreaming && (
                      <AnswerCard
                        content={streamText}
                        sources={_currentSources}
                        images={_currentImages}
                        relatedQuestions={[]}
                        citationMap={_currentCitations}
                        isStreaming
                      />
                    )}
                  </div>
                )}
                <div ref={bottomRef} style={{ height: 1 }} />
              </motion.div>
            )}
          </AnimatePresence>
        </main>

        {!isHomePage && (
          <motion.div initial={{ opacity: 0, y: 20 }} animate={{ opacity: 1, y: 0 }}
            style={{ padding: \'12px 24px 20px\', background: \'var(--bg-primary)\',
                     borderTop: \'1px solid var(--border)\',
                     maxWidth: 760, margin: \'0 auto\', width: \'100%\' }}>
            <SearchBar onSearch={handleSearch} compact placeholder="Ask a follow-up…" />
          </motion.div>
        )}
      </div>
    </div>
  )
}
''', base=FRONTEND)

print("\n" + "="*55)
print("  Document Management System complete!")
print("="*55)
print()
print("Features built:")
print("  1. Guest session — fresh on server restart")
print("  2. Multi-file upload with progress")
print("  3. Document list with [1] [2] [3] numbering")
print("  4. Sequential renumbering on delete")
print("  5. Delete with Qdrant cleanup")
print("  6. Document viewer (PDF text, TXT, DOCX)")
print("  7. [Doc N] citations in LLM responses")
print("  8. Citation map below each answer")
print("  9. Which doc contributed — shown in sources")
print("  10. Processing states: Queued/Parsing/Embedding/Ready")
print()
print("  BONUS: Select specific docs to scope search")
print("  BONUS: Tabs: Documents | History in sidebar")
print()
print("Run:")
print("  uvicorn app.main:app --reload --port 8000")
print("  cd frontend && npm run dev")

# ═══════════════════════════════════════════════════════════
# BACKEND: patch qdrant_service.py — add delete_by_doc
# ═══════════════════════════════════════════════════════════
qdrant_path = os.path.join(BASE, "app", "services", "qdrant_service.py")
if os.path.exists(qdrant_path):
    with open(qdrant_path, "r") as f:
        qs = f.read()
    if "delete_by_doc" not in qs:
        qs += """
    async def delete_by_doc(self, doc_id: str) -> None:
        \"\"\"Remove all text and image points for a document from Qdrant.\"\"\"
        from qdrant_client.models import Filter, FieldCondition, MatchValue
        doc_filter = Filter(must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))])
        try:
            await self._client.delete(
                collection_name=self.text_collection,
                points_selector=doc_filter,
            )
            await self._client.delete(
                collection_name=self.image_collection,
                points_selector=doc_filter,
            )
            logger.info(f"Deleted all Qdrant points for doc_id={doc_id}")
        except Exception as e:
            logger.warning(f"Qdrant delete_by_doc failed: {e}")
"""
        with open(qdrant_path, "w") as f:
            f.write(qs)
        print("  qdrant_service.py: delete_by_doc added")
    else:
        print("  qdrant_service.py: delete_by_doc already present")

# ═══════════════════════════════════════════════════════════
# BACKEND: add CITATIONS to StreamChunkType if missing
# ═══════════════════════════════════════════════════════════
models_path = os.path.join(BASE, "app", "models", "search.py")
if os.path.exists(models_path):
    with open(models_path, "r") as f:
        m = f.read()
    if "CITATIONS" not in m:
        m = m.replace(
            'DONE    = "done"',
            'DONE      = "done"'
        ).replace(
            'ERROR   = "error"',
            'ERROR     = "error"\n    CITATIONS = "citations"'
        )
        with open(models_path, "w") as f:
            f.write(m)
        print("  models/search.py: CITATIONS chunk type added")

print("\nAll patches applied — ready to run!")